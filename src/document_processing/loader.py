from typing import List, Dict, Any
from pathlib import Path
import pypdf
from bs4 import BeautifulSoup
import docx
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging
import re

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentLoader:
    """A class to load and preprocess legal documents from various formats."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the DocumentLoader.
        
        Args:
            chunk_size (int): Size of text chunks for splitting documents
            chunk_overlap (int): Overlap between chunks
        """
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def _preprocess_text(self, text: str) -> str:
        """Clean and preprocess the raw text content."""
        # Convert to lowercase
        text = text.lower()

        # Replace all whitespace sequences (space, tab, newline etc.) with a single space
        text = re.sub(r'\s+', ' ', text).strip()

        # Optional: Remove specific unwanted characters (example)
        # text = re.sub(r'[^a-z0-9\s.,;:\'\"!?-]', '', text)

        return text

    def load_pdf(self, file_path: str) -> List[Document]:
        """Load and process a PDF document."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                raw_text = ""
                for page in pdf_reader.pages:
                    raw_text += page.extract_text() or "" # Add check for None

                # Preprocess the text
                processed_text = self._preprocess_text(raw_text)

                # Split text into chunks
                chunks = self.text_splitter.split_text(processed_text)
                
                # Create Document objects
                documents = [
                    Document(
                        page_content=chunk,
                        metadata={"source": file_path, "type": "pdf"}
                    )
                    for chunk in chunks
                ]
                return documents
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            return []

    def load_txt(self, file_path: str) -> List[Document]:
        """Load and process a text document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                raw_text = file.read()

                # Preprocess the text
                processed_text = self._preprocess_text(raw_text)

                chunks = self.text_splitter.split_text(processed_text)
                documents = [
                    Document(
                        page_content=chunk,
                        metadata={"source": file_path, "type": "txt"}
                    )
                    for chunk in chunks
                ]
                return documents
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {str(e)}")
            return []

    def load_html(self, file_path: str) -> List[Document]:
        """Load and process an HTML document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file, 'html.parser')
                raw_text = soup.get_text()

                # Preprocess the text
                processed_text = self._preprocess_text(raw_text)

                chunks = self.text_splitter.split_text(processed_text)
                documents = [
                    Document(
                        page_content=chunk,
                        metadata={"source": file_path, "type": "html"}
                    )
                    for chunk in chunks
                ]
                return documents
        except Exception as e:
            logger.error(f"Error loading HTML {file_path}: {str(e)}")
            return []

    def load_docx(self, file_path: str) -> List[Document]:
        """Load and process a Word document."""
        try:
            doc = docx.Document(file_path)
            raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Preprocess the text
            processed_text = self._preprocess_text(raw_text)

            chunks = self.text_splitter.split_text(processed_text)
            documents = [
                Document(
                    page_content=chunk,
                    metadata={"source": file_path, "type": "docx"}
                )
                for chunk in chunks
            ]
            return documents
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            return []

    def load_documents(self, directory_path: str) -> List[Document]:
        """
        Load all documents from a directory.
        
        Args:
            directory_path (str): Path to the directory containing documents
            
        Returns:
            List[Document]: List of processed documents
        """
        directory = Path(directory_path)
        all_documents = []
        
        for file_path in directory.glob("**/*"):
            if file_path.is_file():
                if file_path.suffix.lower() == '.pdf':
                    all_documents.extend(self.load_pdf(str(file_path)))
                elif file_path.suffix.lower() == '.txt':
                    all_documents.extend(self.load_txt(str(file_path)))
                elif file_path.suffix.lower() == '.html':
                    all_documents.extend(self.load_html(str(file_path)))
                elif file_path.suffix.lower() == '.docx':
                    all_documents.extend(self.load_docx(str(file_path)))
        
        return all_documents 